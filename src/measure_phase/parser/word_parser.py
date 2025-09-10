
"""
Word Document Parser (.docx)
============================

Parser for Microsoft Word documents using python-docx.
Extracts text, tables, images, and document structure.
"""

from typing import Dict, Any, List
import logging

from .base_parser import BaseParser

logger = logging.getLogger(__name__)

class WordParser(BaseParser):
    """Parser for Word .docx files"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.docx']
        
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Parse Word document and extract content
        
        Args:
            file_path: Path to .docx file
            
        Returns:
            Dictionary with document analysis
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        "text": para.text,
                        "style": para.style.name if para.style else "Normal"
                    })
            
            # Extract tables
            tables = []
            for table_idx, table in enumerate(doc.tables):
                table_data = self._extract_table_data(table, table_idx)
                tables.append(table_data)
            
            # Extract document properties
            properties = self._extract_properties(doc)
            
            # Analyze document structure
            structure = self._analyze_structure(paragraphs)
            
            return {
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
                "word_count": sum(len(p["text"].split()) for p in paragraphs),
                "character_count": sum(len(p["text"]) for p in paragraphs),
                "paragraphs": paragraphs,
                "tables": tables,
                "properties": properties,
                "structure": structure,
                "document_type": self._identify_document_type(paragraphs)
            }
            
        except ImportError:
            logger.error("python-docx library not installed. Install with: pip install python-docx")
            return {"error": "python-docx library not available"}
        except Exception as e:
            logger.error(f"Error parsing Word document {file_path}: {e}")
            return {"error": str(e)}
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract plain text from Word document
        
        Args:
            file_path: Path to .docx file
            
        Returns:
            Plain text content
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_content = []
            
            # Extract paragraph text
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extract table text
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text_content.append(" | ".join(row_text))
            
            return "\n".join(text_content)
            
        except ImportError:
            return "Error: python-docx library not available"
        except Exception as e:
            logger.error(f"Error extracting text from Word document {file_path}: {e}")
            return f"Error: {e}"
    
    def _extract_table_data(self, table, table_idx: int) -> Dict[str, Any]:
        """Extract data from a table"""
        rows = []
        
        for row_idx, row in enumerate(table.rows):
            cells = []
            for cell in row.cells:
                cells.append(cell.text.strip())
            rows.append(cells)
        
        return {
            "table_index": table_idx,
            "row_count": len(rows),
            "column_count": len(rows[0]) if rows else 0,
            "data": rows
        }
    
    def _extract_properties(self, doc) -> Dict[str, Any]:
        """Extract document properties"""
        properties = {}
        
        try:
            core_props = doc.core_properties
            properties.update({
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "keywords": core_props.keywords or ""
            })
        except Exception as e:
            logger.debug(f"Could not extract document properties: {e}")
        
        return properties
    
    def _analyze_structure(self, paragraphs: List[Dict]) -> Dict[str, Any]:
        """Analyze document structure"""
        structure = {
            "headings": [],
            "normal_paragraphs": 0,
            "styles_used": set()
        }
        
        for para in paragraphs:
            style = para.get("style", "Normal")
            structure["styles_used"].add(style)
            
            if "Heading" in style:
                structure["headings"].append({
                    "text": para["text"],
                    "level": style
                })
            else:
                structure["normal_paragraphs"] += 1
        
        structure["styles_used"] = list(structure["styles_used"])
        return structure
    
    def _identify_document_type(self, paragraphs: List[Dict]) -> str:
        """Identify document type based on content"""
        text_content = " ".join([p["text"] for p in paragraphs]).lower()
        
        # Check for common document types
        if any(word in text_content for word in ['requirements', 'specification', 'spec']):
            return "Requirements Document"
        elif any(word in text_content for word in ['manual', 'guide', 'instructions']):
            return "Manual/Guide"
        elif any(word in text_content for word in ['report', 'analysis', 'findings']):
            return "Report"
        elif any(word in text_content for word in ['proposal', 'plan', 'strategy']):
            return "Proposal/Plan"
        elif any(word in text_content for word in ['meeting', 'minutes', 'agenda']):
            return "Meeting Document"
        else:
            return "General Document"
